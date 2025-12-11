"""
文档加载服务
负责从本地文件系统加载 Markdown/TXT/Word 笔记并提取元数据
"""
import os
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any, Optional
import frontmatter
from llama_index.core import Document
from llama_index.core.readers import SimpleDirectoryReader

from app.config import settings


class DocumentLoader:
    """文档加载器 - 支持 Markdown、TXT、Word"""
    
    # 支持的文件格式
    MARKDOWN_EXTENSIONS = [".md", ".markdown"]
    TEXT_EXTENSIONS = [".txt"]
    WORD_EXTENSIONS = [".docx"]
    
    SUPPORTED_EXTENSIONS = MARKDOWN_EXTENSIONS + TEXT_EXTENSIONS + WORD_EXTENSIONS
    
    def __init__(self, directories: Optional[List[Path]] = None):
        """
        初始化文档加载器
        
        Args:
            directories: 笔记目录列表，默认从配置读取
        """
        self.directories = directories or settings.notes_paths
    
    def load_all_documents(self) -> List[Document]:
        """
        加载所有目录下的文档
        
        Returns:
            LlamaIndex Document 列表
        """
        all_documents = []
        
        for directory in self.directories:
            if directory.exists():
                documents = self._load_from_directory(directory)
                all_documents.extend(documents)
        
        return all_documents
    
    def _load_from_directory(self, directory: Path) -> List[Document]:
        """
        从单个目录加载文档
        
        Args:
            directory: 目录路径
            
        Returns:
            Document 列表
        """
        documents = []
        
        for file_path in directory.rglob("*"):
            if file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    doc = self._load_single_file(file_path)
                    if doc:
                        documents.append(doc)
                except Exception as e:
                    print(f"Error loading {file_path}: {e}")
        
        return documents
    
    def _load_single_file(self, file_path: Path) -> Optional[Document]:
        """
        加载单个文件并提取元数据
        
        Args:
            file_path: 文件路径
            
        Returns:
            LlamaIndex Document 或 None
        """
        try:
            suffix = file_path.suffix.lower()
            
            # 根据文件类型选择加载方式
            if suffix in self.WORD_EXTENSIONS:
                content = self._load_word_file(file_path)
            else:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            
            if not content:
                return None
            
            # 提取元数据
            metadata = self._extract_metadata(file_path, content)
            
            # 如果是 markdown，解析 frontmatter
            if suffix in self.MARKDOWN_EXTENSIONS:
                post = frontmatter.loads(content)
                # 合并 frontmatter 中的元数据
                metadata.update(post.metadata)
                content = post.content
            
            return Document(
                text=content,
                metadata=metadata,
                doc_id=str(file_path)
            )
            
        except Exception as e:
            print(f"Failed to load {file_path}: {e}")
            return None
    
    def _load_word_file(self, file_path: Path) -> Optional[str]:
        """
        加载 Word 文档 (.docx)
        
        Args:
            file_path: Word 文件路径
            
        Returns:
            文档文本内容
        """
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(str(file_path))
            
            # 提取所有段落文本
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            
            return "\n\n".join(paragraphs)
            
        except ImportError:
            print("⚠️ python-docx not installed. Run: pip install python-docx")
            return None
        except Exception as e:
            print(f"Failed to load Word file {file_path}: {e}")
            return None
    
    def _extract_metadata(self, file_path: Path, content: str) -> Dict[str, Any]:
        """
        提取文件元数据
        
        Args:
            file_path: 文件路径
            content: 文件内容
            
        Returns:
            元数据字典
        """
        stat = file_path.stat()
        
        # 从文件名提取标题
        title = file_path.stem
        
        # 尝试从内容提取第一个标题
        lines = content.split("\n")
        for line in lines:
            if line.startswith("# "):
                title = line[2:].strip()
                break
        
        return {
            "file_name": file_path.name,
            "file_path": str(file_path.absolute()),
            "title": title,
            "extension": file_path.suffix,
            "file_type": self._get_file_type(file_path.suffix.lower()),
            "created_at": datetime.fromtimestamp(stat.st_birthtime).isoformat() if hasattr(stat, 'st_birthtime') else datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "updated_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
            "file_size": stat.st_size,
            "word_count": len(content.split()),
        }
    
    def _get_file_type(self, suffix: str) -> str:
        """获取文件类型标签"""
        if suffix in self.MARKDOWN_EXTENSIONS:
            return "markdown"
        elif suffix in self.WORD_EXTENSIONS:
            return "word"
        else:
            return "text"
    
    def load_single_document(self, file_path: str) -> Optional[Document]:
        """
        加载单个文档（用于增量更新）
        
        Args:
            file_path: 文件路径字符串
            
        Returns:
            Document 或 None
        """
        path = Path(file_path)
        if path.exists() and path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
            return self._load_single_file(path)
        return None


# 全局实例
document_loader = DocumentLoader()
